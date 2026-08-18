import os
import re
import pandas as pd
import openpyxl
from docx import Document

DATA_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), 'data'))

UNILOG_252_HEADERS = [
    "MFR URL", "Ref URL 1", "Ref URL 2", "Ref URL 3", "Ref URL 4", "Ref URL 5", "PART_NUMBER", "Dept", "Class", "Fine",
    "SKU - MY_PART_NUMBER", "Mfg_Part_Num", "Part_Desc", "E1_Brand", "Unilog_Brand", "DIB_Brand", "Part_Manuf",
    "MANUFACTURER_NAME", "BRAND_NAME", "TRADE_NAME", "MANUFACTURER_PART_NUMBER", "ALTERNATE_PART_NUMBER", "Classpath",
    "MOBILE_DESC", "INVOICE_DESC", "SHORT_DESC", "LONG_DESC1", "RETAIL_DESC", "MARKETING_DESCRIPTION",
    "ITEM_FEATURES_1", "ITEM_FEATURES_2", "ITEM_FEATURES_3", "ITEM_FEATURES_4", "ITEM_FEATURES_5",
    "ITEM_FEATURES_6", "ITEM_FEATURES_7", "ITEM_FEATURES_8", "ITEM_FEATURES_9", "ITEM_FEATURES_10",
    "ITEM_FEATURES_11", "ITEM_FEATURES_12", "ITEM_FEATURES_13", "ITEM_FEATURES_14", "ITEM_FEATURES_15",
    "ITEM_FEATURES_16", "ITEM_FEATURES_17", "ITEM_FEATURES_18", "ITEM_FEATURES_19", "ITEM_FEATURES_20",
    "With", "Standard/Approvals", "Prop 65", "Application", "Includes", "Product Name"
]
for i in range(1, 51):
    UNILOG_252_HEADERS.extend([f"ATTRIBUTE_LABEL {i}", f"ATTRIBUTE_VALUE {i}", f"ATTRIBUTE_UOM {i}"])

UNILOG_252_HEADERS.extend([
    "UPC", "EAN", "GTIN", "UNSPSC", "Warranty", "List Price", "Selling Qty", "Selling UOM",
    "Standard Packaging Information", "LENGTH", "LENGTH_UOM", "HEIGHT", "HEIGHT_UOM", "WIDTH", "WIDTH_UOM",
    "WEIGHT", "WEIGHT_UOM", "VOLUME", "VOLUME_UOM", "Product Image", "Alternate Image 1", "Alternate Image 2",
    "Alternate Image 3", "Alternate Image 4", "SDS", "SDS_1", "Warranty Information", "Catalog",
    "Specification Sheet", "Instruction/Installation Manual", "Service Manual", "Owners/User Manual",
    "Line Drawing", "MTR", "RoHS", "Full Engineering Drawing", "Energy Star Guide", "Technical Bulletin",
    "Submittal", "Compatibility Chart", "Size Chart", "Product Label/Insert", "Video Link", "Video Link 1",
    "Country Of Origin", "Discontinued", "Actual Image (Yes/No)"
])

# Ensure exact 252 length
assert len(UNILOG_252_HEADERS) == 252, f"Expected 252 headers, got {len(UNILOG_252_HEADERS)}"


def setup_data_files():
    os.makedirs(DATA_DIR, exist_ok=True)

    # 1. Decimal_Fraction.xlsx (4 stacked Fraction|Decimal column pairs)
    dec_frac_path = os.path.join(DATA_DIR, 'Decimal_Fraction.xlsx')
    if not os.path.exists(dec_frac_path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["Fraction", "Decimal", "Fraction", "Decimal", "Fraction", "Decimal", "Fraction", "Decimal"])
        
        pairs = [
            ("1/16", 0.0625), ("1/8", 0.125), ("3/16", 0.1875), ("1/4", 0.25),
            ("5/16", 0.3125), ("3/8", 0.375), ("7/16", 0.4375), ("1/2", 0.5),
            ("9/16", 0.5625), ("5/8", 0.625), ("11/16", 0.6875), ("3/4", 0.75),
            ("13/16", 0.8125), ("7/8", 0.875), ("15/16", 0.9375), ("1", 1.0)
        ]
        
        for i in range(0, len(pairs), 4):
            row = []
            for j in range(4):
                if i + j < len(pairs):
                    row.extend([pairs[i+j][0], pairs[i+j][1]])
                else:
                    row.extend(["", ""])
            ws.append(row)
        wb.save(dec_frac_path)

    # 2. Unilog_Master_UOM_Standards_Abbreviations_and_Terms.xlsx
    uom_path = os.path.join(DATA_DIR, 'Unilog_Master_UOM_Standards_Abbreviations_and_Terms.xlsx')
    if not os.path.exists(uom_path):
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["Term / Variant", "Canonical Abbreviation", "Description"])
        uom_data = [
            ("inch", "in", "Length"), ("inches", "in", "Length"), ("in.", "in", "Length"), ('"', "in", "Length"),
            ("foot", "ft", "Length"), ("feet", "ft", "Length"), ("ft.", "ft", "Length"), ("'", "ft", "Length"),
            ("millimeter", "mm", "Length"), ("millimeters", "mm", "Length"), ("mm.", "mm", "Length"),
            ("centimeter", "cm", "Length"), ("centimeters", "cm", "Length"), ("cm.", "cm", "Length"),
            ("meter", "m", "Length"), ("meters", "m", "Length"), ("m.", "m", "Length"),
            ("pound", "lb", "Weight"), ("pounds", "lb", "Weight"), ("lb.", "lb", "Weight"), ("lbs", "lb", "Weight"),
            ("ounce", "oz", "Weight"), ("ounces", "oz", "Weight"), ("oz.", "oz", "Weight"),
            ("gallon", "gal", "Volume"), ("gallons", "gal", "Volume"), ("gal.", "gal", "Volume"),
            ("gallons per minute", "gpm", "Flow Rate"), ("GPM", "gpm", "Flow Rate"), ("gpm", "gpm", "Flow Rate"),
            ("gallons per day", "gpd", "Flow Rate"), ("GPD", "gpd", "Flow Rate"),
            ("volt", "V", "Electrical"), ("volts", "V", "Electrical"), ("V", "V", "Electrical"), ("v", "V", "Electrical"),
            ("ampere", "A", "Electrical"), ("amp", "A", "Electrical"), ("amps", "A", "Electrical"), ("A", "A", "Electrical"),
            ("watt", "W", "Electrical"), ("watts", "W", "Electrical"), ("W", "W", "Electrical"),
            ("psi", "psi", "Pressure"), ("PSI", "psi", "Pressure"), ("pounds per square inch", "psi", "Pressure"),
            ("degree fahrenheit", "deg F", "Temperature"), ("deg F", "deg F", "Temperature"), ("°F", "deg F", "Temperature"),
            ("rpm", "rpm", "Speed"), ("RPM", "rpm", "Speed")
        ]
        for row in uom_data:
            ws1.append(row)

        ws2 = wb.create_sheet(title="Sheet2")
        ws2.append(["Rule Name", "House Style Rule Description"])
        ws2.append(["Spacing", "Always put a space between number and unit, e.g. 24 in not 24in"])
        ws2.append(["Fractions", "Use fractional representations for decimals as defined in house lookup"])
        ws2.append(["Casing", "Follow canonical UOM casing (e.g., V for Volts, A for Amps, gpm for Flow Rate)"])
        wb.save(uom_path)

    # 3. UniCat_Manufacturer_and_Brand_List.xlsx
    mfr_path = os.path.join(DATA_DIR, 'UniCat_Manufacturer_and_Brand_List.xlsx')
    if not os.path.exists(mfr_path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Approved List"
        ws.append(["MANUFACTURER_NAME", "MANUFACTURER_CODE", "BRAND_NAME", "BRAND_CODE"])
        mfr_list = [
            ("Freud Inc", "2435", "DIABLO®", "DIABLO"),
            ("3M Co", "5293", "3M®", "3M"),
            ("Mirka Abrasives Inc", "MIRUS", "MIRKA®", "MIRKA"),
            ("Milwaukee Accessory", "4031", "MILWAUKEE®", "MILW"),
            ("Appliance Dealers Cooperative", "APPDE", "FRIGIDAIRE®", "FRIG"),
            ("Appliance Dealers Cooperative", "APPDE", "Whirlpool®", "WHIRL"),
            ("Appliance Dealers Cooperative", "APPDE", "GE®", "GE"),
            ("Appliance Dealers Cooperative", "APPDE", "LG®", "LG"),
            ("Appliance Dealers Cooperative", "APPDE", "KitchenAid®", "KITCH"),
            ("Appliance Dealers Cooperative", "APPDE", "Speed Queen®", "SPEED"),
            ("Appliance Dealers Cooperative", "APPDE", "Café™", "CAFE"),
            ("Moen Incorporated", "MOEN", "MOEN®", "MOEN"),
            ("Delta Faucet Company", "DELTA", "DELTA®", "DELTA"),
            ("Kohler Co.", "KOHLER", "KOHLER®", "KOHLER"),
            ("American Standard", "AMSTD", "American Standard®", "AMSTD"),
            ("Pfister", "PFIST", "Pfister®", "PFIST"),
            ("Nibco Inc", "NIBCO", "NIBCO®", "NIBCO"),
            ("Charlotte Pipe and Foundry", "CHARL", "CHARLOTTE PIPE®", "CHARL"),
            ("Mueller Streamline", "MUELL", "STREAMLINE®", "STREAM"),
            ("Viega LLC", "VIEGA", "VIEGA®", "VIEGA"),
            ("SharkBite Plumbing Solutions", "SHARK", "SHARKBITE®", "SHARK"),
            ("Black & Decker", "2585", "DEWALT®", "DEWALT"),
            ("Kreg Tool Company", "KRETO", "KREG®", "KREG"),
            ("Makita Usa Inc", "5142", "MAKITA®", "MAKITA"),
            ("Satco Prod Inc", "5573", "SATCO®", "SATCO"),
            ("Leviton Mfg Co", "4927", "LEVITON®", "LEVITON"),
            ("Kichler Lighting", "KICLI", "KICHLER®", "KICHLER")
        ]
        for row in mfr_list:
            ws.append(row)
        wb.save(mfr_path)

    # 4. FAUCETS_LOV.xlsx
    faucets_path = os.path.join(DATA_DIR, 'FAUCETS_LOV.xlsx')
    if not os.path.exists(faucets_path):
        wb = openpyxl.Workbook()
        ws_sum = wb.active
        ws_sum.title = "Summary"
        ws_sum.append(["Classpath", "UNSPSC", "Category Name"])
        faucets_classpaths = [
            ("Plumbing>Faucets>Bathroom Sink Faucets", "30181702", "Bathroom Sink Faucets"),
            ("Plumbing>Faucets>Kitchen Faucets", "30181703", "Kitchen Faucets"),
            ("Plumbing>Faucets>Tub & Shower Faucets", "30181704", "Tub & Shower Faucets"),
            ("Plumbing>Faucets>Bar Faucets", "30181705", "Bar Faucets"),
            ("Plumbing>Faucets>Utility Faucets", "30181706", "Utility Faucets")
        ]
        for row in faucets_classpaths:
            ws_sum.append(row)

        ws_det = wb.create_sheet(title="Attribute Detail")
        ws_det.append(["Classpath", "Attribute Label", "Attribute Order", "Filtering Y/N", "Allowed Values", "Normalized Values", "UOM"])
        faucet_attrs = [
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Faucet Handle Type", 1, "Y", "Lever, Cross, Knob, Touchless, Wrist Blade", "Lever, Cross, Knob, Touchless, Wrist Blade", ""),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Number of Faucet Handles", 2, "Y", "1, 2, 3", "1, 2, 3", ""),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Faucet Hole Center", 3, "Y", "4 in, 8 in, Single Hole", "4 in, 8 in, Single Hole", "in"),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Flow Rate", 4, "Y", "0.5 gpm, 1.0 gpm, 1.2 gpm, 1.5 gpm", "0.5 gpm, 1.0 gpm, 1.2 gpm, 1.5 gpm", "gpm"),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Finish/Color", 5, "Y", "Chrome, Brushed Nickel, Matte Black, Oil Rubbed Bronze, Polished Brass", "Chrome, Brushed Nickel, Matte Black, Oil Rubbed Bronze, Polished Brass", ""),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Material", 6, "N", "Brass, Stainless Steel, Zinc", "Brass, Stainless Steel, Zinc", ""),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Spout Reach", 7, "N", "4-1/2 in, 5 in, 5-1/2 in, 6 in", "4-1/2 in, 5 in, 5-1/2 in, 6 in", "in"),
            ("Plumbing>Faucets>Kitchen Faucets", "Faucet Type", 1, "Y", "Pull-Down, Pull-Out, Commercial Style, Standard Spout", "Pull-Down, Pull-Out, Commercial Style, Standard Spout", ""),
            ("Plumbing>Faucets>Kitchen Faucets", "Number of Faucet Handles", 2, "Y", "1, 2", "1, 2", ""),
            ("Plumbing>Faucets>Kitchen Faucets", "Flow Rate", 3, "Y", "1.5 gpm, 1.8 gpm, 2.2 gpm", "1.5 gpm, 1.8 gpm, 2.2 gpm", "gpm"),
            ("Plumbing>Faucets>Kitchen Faucets", "Finish/Color", 4, "Y", "Chrome, Stainless Steel, Matte Black, Brushed Gold", "Chrome, Stainless Steel, Matte Black, Brushed Gold", ""),
            ("Plumbing>Faucets>Kitchen Faucets", "Number of Faucet Holes", 5, "Y", "1, 2, 3, 4", "1, 2, 3, 4", "")
        ]
        for row in faucet_attrs:
            ws_det.append(row)
        wb.save(faucets_path)

    # 5. Fittings_LOV.xlsx
    fittings_path = os.path.join(DATA_DIR, 'Fittings_LOV.xlsx')
    if not os.path.exists(fittings_path):
        wb = openpyxl.Workbook()
        ws_sum = wb.active
        ws_sum.title = "Summary"
        ws_sum.append(["Classpath", "UNSPSC", "Category Name"])
        fittings_classpaths = [
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "40172600", "Pipe Fittings"),
            ("Plumbing>Pipe, Tube & Hose Fittings>Tube Fittings", "40172601", "Tube Fittings"),
            ("Plumbing>Pipe, Tube & Hose Fittings>Hose Fittings", "40172602", "Hose Fittings"),
            ("Plumbing>Pipe, Tube & Hose Fittings>Nipples", "40172603", "Nipples")
        ]
        for row in fittings_classpaths:
            ws_sum.append(row)

        ws_det = wb.create_sheet(title="Attribute Detail")
        ws_det.append(["Classpath", "Attribute Label", "Attribute Order", "Filtering Y/N", "Allowed Values", "Normalized Values", "UOM"])
        fitting_attrs = [
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Fitting Fitting/Connector Type", 1, "Y", "Elbow, Tee, Coupling, Adapter, Cap, Plug, Union, Reducer", "Elbow, Tee, Coupling, Adapter, Cap, Plug, Union, Reducer", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Pipe Size", 2, "Y", "1/4 in, 3/8 in, 1/2 in, 3/4 in, 1 in, 1-1/4 in, 1-1/2 in, 2 in", "1/4 in, 3/8 in, 1/2 in, 3/4 in, 1 in, 1-1/4 in, 1-1/2 in, 2 in", "in"),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Connection Type", 3, "Y", "FNPT, MNPT, Push-to-Connect, Press, Sweat, Slip, Threaded", "FNPT, MNPT, Push-to-Connect, Press, Sweat, Slip, Threaded", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Fitting Material", 4, "Y", "Brass, Copper, PVC, CPVC, Stainless Steel, Carbon Steel, Ductile Iron", "Brass, Copper, PVC, CPVC, Stainless Steel, Carbon Steel, Ductile Iron", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Fitting Schedule/Class", 5, "N", "Schedule 40, Schedule 80, Class 150, Class 300", "Schedule 40, Schedule 80, Class 150, Class 300", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Nipples", "Pipe Size", 1, "Y", "1/4 in, 3/8 in, 1/2 in, 3/4 in, 1 in, 1-1/4 in, 1-1/2 in, 2 in", "1/4 in, 3/8 in, 1/2 in, 3/4 in, 1 in, 1-1/4 in, 1-1/2 in, 2 in", "in"),
            ("Plumbing>Pipe, Tube & Hose Fittings>Nipples", "Length", 2, "Y", "Close, 1-1/2 in, 2 in, 2-1/2 in, 3 in, 3-1/2 in, 4 in, 6 in, 8 in, 10 in, 12 in", "Close, 1-1/2 in, 2 in, 2-1/2 in, 3 in, 3-1/2 in, 4 in, 6 in, 8 in, 10 in, 12 in", "in"),
            ("Plumbing>Pipe, Tube & Hose Fittings>Nipples", "Connection Type", 3, "Y", "MNPT x MNPT", "MNPT x MNPT", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Nipples", "Fitting Material", 4, "Y", "Black Steel, Galvanized Steel, Brass, Stainless Steel", "Black Steel, Galvanized Steel, Brass, Stainless Steel", "")
        ]
        for row in fitting_attrs:
            ws_det.append(row)
        wb.save(fittings_path)

    # 6. Unicat_Lov_v1_0_Updated_With_Remarks.xlsx
    unicat_lov_path = os.path.join(DATA_DIR, 'Unicat_Lov_v1_0_Updated_With_Remarks.xlsx')
    if not os.path.exists(unicat_lov_path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "LOV Master"
        ws.append(["Classpath", "Leaf Node", "Filtering Y/N", "Attribute Label", "Attribute Values", "Normalized Label", "Normalized Values", "Guidelines", "Remarks"])
        
        lov_rows = [
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Bathroom Sink Faucets", "Y", "Faucet Handle Type", "Lever, Cross, Knob, Touchless, Wrist Blade", "Faucet Handle Type", "Lever, Cross, Knob, Touchless, Wrist Blade", "Follow house rules", ""),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Bathroom Sink Faucets", "Y", "Number of Faucet Handles", "1, 2, 3", "Number of Faucet Handles", "1, 2, 3", "Numeric count", ""),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Bathroom Sink Faucets", "Y", "Faucet Hole Center", "4 in, 8 in, Single Hole", "Faucet Hole Center", "4 in, 8 in, Single Hole", "Standard centerset", ""),
            ("Plumbing>Faucets>Bathroom Sink Faucets", "Bathroom Sink Faucets", "Y", "Flow Rate", "0.5 gpm, 1.0 gpm, 1.2 gpm, 1.5 gpm", "Flow Rate", "0.5 gpm, 1.0 gpm, 1.2 gpm, 1.5 gpm", "WaterSense flow rate", ""),
            ("Plumbing>Faucets>Kitchen Faucets", "Kitchen Faucets", "Y", "Faucet Type", "Pull-Down, Pull-Out, Commercial Style, Standard Spout", "Faucet Type", "Pull-Down, Pull-Out, Commercial Style, Standard Spout", "Kitchen faucet style", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Pipe Fittings", "Y", "Fitting Fitting/Connector Type", "Elbow, Tee, Coupling, Adapter, Cap, Plug, Union, Reducer", "Fitting Fitting/Connector Type", "Elbow, Tee, Coupling, Adapter, Cap, Plug, Union, Reducer", "Fitting shape", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Pipe Fittings", "Y", "Pipe Size", "1/4 in, 3/8 in, 1/2 in, 3/4 in, 1 in, 1-1/4 in, 1-1/2 in, 2 in", "Pipe Size", "1/4 in, 3/8 in, 1/2 in, 3/4 in, 1 in, 1-1/4 in, 1-1/2 in, 2 in", "Nominal size", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Pipe Fittings", "Y", "Connection Type", "FNPT, MNPT, Push-to-Connect, Press, Sweat, Slip, Threaded", "Connection Type", "FNPT, MNPT, Push-to-Connect, Press, Sweat, Slip, Threaded", "End connection", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Pipe Fittings", "Y", "Fitting Material", "Brass, Copper, PVC, CPVC, Stainless Steel, Carbon Steel, Ductile Iron", "Fitting Material", "Brass, Copper, PVC, CPVC, Stainless Steel, Carbon Steel, Ductile Iron", "Body material", ""),
            ("Plumbing>Pipe, Tube & Hose Fittings>Nipples", "Nipples", "Y", "Length", "Close, 1-1/2 in, 2 in, 2-1/2 in, 3 in, 3-1/2 in, 4 in, 6 in, 8 in, 10 in, 12 in", "Length", "Close, 1-1/2 in, 2 in, 2-1/2 in, 3 in, 3-1/2 in, 4 in, 6 in, 8 in, 10 in, 12 in", "Overall length", ""),
            ("Abrasives>Sanding Belts & Discs>Sanding Belts", "Sanding Belts", "Y", "Grit", "60, 80, 120, 150, 180, 220, 320", "Grit", "60, 80, 120, 150, 180, 220, 320", "Abrasive grit", ""),
            ("Appliances & Consumer Electronics>Kitchen Appliances>Built-In Dishwashers", "Built-In Dishwashers", "Y", "Number of Wash Cycles", "5, 6, 7", "Number of Wash Cycles", "5, 6, 7", "Dishwasher cycles", "")
        ]
        for row in lov_rows:
            ws.append(row)
        wb.save(unicat_lov_path)

    # 7. UNILOG_INTERNAL_CONTENT_GUIDELINES.docx
    doc_path = os.path.join(DATA_DIR, 'UNILOG_INTERNAL_CONTENT_GUIDELINES.docx')
    if not os.path.exists(doc_path):
        doc = Document()
        doc.add_heading('UNILOG INTERNAL CONTENT GUIDELINES', 0)
        doc.add_heading('Description Field Formulas', level=1)
        doc.add_paragraph('INVOICE_DESC: Maximum 40 characters. All UPPERCASE. Formula: BRAND + MPN + SHORT ITEM TYPE.')
        doc.add_paragraph('MOBILE_DESC: 60-80 characters. Formula: BRAND + SERIES + MPN + ITEM TYPE + KEY ATTRIBUTES.')
        doc.add_paragraph('SHORT_DESC: Formula: BRAND® + SERIES + MPN + ITEM TYPE + KEY SPECIFICATIONS.')
        doc.add_paragraph('LONG_DESC1: Full detailed prose with all extracted attributes, UOMs, and features.')
        doc.add_heading('Character Limits & Casing', level=1)
        doc.add_paragraph('INVOICE_DESC <= 40 chars. MOBILE_DESC 60-80 chars.')
        doc.add_paragraph('Placeholders like "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --" MUST be converted to NULL.')
        doc.save(doc_path)

    # 8. Reference_Documents_Summary.xlsx
    summary_path = os.path.join(DATA_DIR, 'Reference_Documents_Summary.xlsx')
    if not os.path.exists(summary_path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Summary"
        ws.append(["File Name", "Description", "Primary Sheet", "Usage"])
        ws.append(["Sample-1000_Items.xlsx", "1,000 Bulk Raw Input Items", "Sheet1", "Scale Test Input"])
        ws.append(["Unilog-Sample_200_Items-Input-vs-Output.xlsx", "Ground Truth Input & 252-col Delivery Format", "Delivery Format", "Ground Truth & Header Source"])
        ws.append(["UniCat_Manufacturer_and_Brand_List.xlsx", "Approved Manufacturers & Brands", "Approved List", "Manufacturer/Brand Resolver"])
        ws.append(["Unicat_Lov_v1_0_Updated_With_Remarks.xlsx", "Master LOV Vocabulary Specs", "LOV Master", "Classifier & Attribute Extractor"])
        ws.append(["FAUCETS_LOV.xlsx", "Faucets Deep Spec", "Summary / Attribute Detail", "Faucets Reference Implementation"])
        ws.append(["Fittings_LOV.xlsx", "Fittings Deep Spec", "Summary / Attribute Detail", "Fittings Reference Implementation"])
        ws.append(["Unilog_Master_UOM_Standards_Abbreviations_and_Terms.xlsx", "UOM Standards & Rules", "Sheet1 / Sheet2", "UOM Normalization"])
        ws.append(["Decimal_Fraction.xlsx", "Decimal to Fraction Lookup", "Sheet1", "Decimal/Fraction Normalization"])
        wb.save(summary_path)

    # 9. Unilog-Sample_200_Items-Input-vs-Output.xlsx & Sample-1000_Items.xlsx
    gt_path = os.path.join(DATA_DIR, 'Unilog-Sample_200_Items-Input-vs-Output.xlsx')
    sample_1000_path = os.path.join(DATA_DIR, 'Sample-1000_Items.xlsx')

    if not os.path.exists(gt_path) or not os.path.exists(sample_1000_path):
        # Create GT workbook
        wb_gt = openpyxl.Workbook()
        ws_in = wb_gt.active
        ws_in.title = "Input"
        ws_in.append(["Mfg_Part_Num", "Part_Desc", "E1_Brand", "Unilog_Brand", "DIB_Brand", "Part_Manuf"])
        
        ws_out = wb_gt.create_sheet(title="Delivery Format")
        ws_out.append(UNILOG_252_HEADERS)

        # Faucet and Fitting sample items for Ground Truth + General sample items
        gt_sample_items = [
            # Faucets
            ("6702-000", "Moen 6702-000 Genta Single Handle Bathroom Sink Faucet Chrome 1.2 gpm 4in", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Moen Incorporated (MOEN)",
             "Plumbing>Faucets>Bathroom Sink Faucets", "Moen Incorporated", "MOEN®", "6702-000", "MOEN 6702-000 GENTA BATH FAUCET CHR 1.2GPM",
             "MOEN® Genta 6702-000 Single Handle Bathroom Sink Faucet, Chrome",
             "MOEN® Genta 6702-000 Single Handle Bathroom Sink Faucet, 1.2 gpm, Chrome, 4 in Centers",
             [("Faucet Handle Type", "Lever", ""), ("Number of Faucet Handles", "1", ""), ("Faucet Hole Center", "4 in", "in"), ("Flow Rate", "1.2 gpm", "gpm"), ("Finish/Color", "Chrome", "")]),
            
            ("559HA-DST", "Delta 559HA-DST Trinsic Single Handle Bathroom Sink Faucet Matte Black 1.2 gpm", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Delta Faucet Company (DELTA)",
             "Plumbing>Faucets>Bathroom Sink Faucets", "Delta Faucet Company", "DELTA®", "559HA-DST", "DELTA 559HA-DST TRINSIC BATH FAUCET BLK",
             "DELTA® Trinsic 559HA-DST Single Handle Bathroom Sink Faucet, Matte Black",
             "DELTA® Trinsic 559HA-DST Single Handle Bathroom Sink Faucet, 1.2 gpm, Matte Black",
             [("Faucet Handle Type", "Lever", ""), ("Number of Faucet Handles", "1", ""), ("Flow Rate", "1.2 gpm", "gpm"), ("Finish/Color", "Matte Black", "")]),

            ("K-596-VS", "Kohler K-596-VS Simplice Pull-Down Kitchen Faucet Vibrant Stainless 1.5 gpm", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Kohler Co. (KOHLER)",
             "Plumbing>Faucets>Kitchen Faucets", "Kohler Co.", "KOHLER®", "K-596-VS", "KOHLER K-596-VS SIMPLICE KITCHEN FAUCET SS",
             "KOHLER® Simplice K-596-VS Pull-Down Kitchen Faucet, Vibrant Stainless",
             "KOHLER® Simplice K-596-VS Pull-Down Kitchen Faucet, 1.5 gpm, Vibrant Stainless",
             [("Faucet Type", "Pull-Down", ""), ("Number of Faucet Handles", "1", ""), ("Flow Rate", "1.5 gpm", "gpm"), ("Finish/Color", "Stainless Steel", "")]),

            # Fittings
            ("607-1/2", "Nibco 607 1/2 1/2 in Copper Wrot Pressure 90 Deg Elbow", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Nibco Inc (NIBCO)",
             "Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "Nibco Inc", "NIBCO®", "607-1/2", "NIBCO 607 1/2IN COPPER 90DEG ELBOW",
             "NIBCO® 607 1/2 in 90 Deg Copper Pressure Elbow",
             "NIBCO® 607 1/2 in 90 Deg Wrot Copper Pressure Elbow Fitting",
             [("Fitting Fitting/Connector Type", "Elbow", ""), ("Pipe Size", "1/2 in", "in"), ("Connection Type", "Sweat", ""), ("Fitting Material", "Copper", "")]),

            ("U008LF", "SharkBite U008LF 1/2 in Push-to-Connect Brass Coupling", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "SharkBite Plumbing Solutions (SHARK)",
             "Plumbing>Pipe, Tube & Hose Fittings>Pipe Fittings", "SharkBite Plumbing Solutions", "SHARKBITE®", "U008LF", "SHARKBITE U008LF 1/2IN PUSH COUPLING BRASS",
             "SHARKBITE® U008LF 1/2 in Push-to-Connect Straight Coupling",
             "SHARKBITE® U008LF 1/2 in Push-to-Connect Lead-Free Brass Coupling",
             [("Fitting Fitting/Connector Type", "Coupling", ""), ("Pipe Size", "1/2 in", "in"), ("Connection Type", "Push-to-Connect", ""), ("Fitting Material", "Brass", "")]),

            ("NIP-050-CLOSE", "Charlotte Pipe 1/2 in x Close Black Steel Nipple MNPT", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Charlotte Pipe and Foundry (CHARL)",
             "Plumbing>Pipe, Tube & Hose Fittings>Nipples", "Charlotte Pipe and Foundry", "CHARLOTTE PIPE®", "NIP-050-CLOSE", "CHARLOTTE 1/2IN CLOSE BLK STEEL NIPPLE",
             "CHARLOTTE PIPE® 1/2 in x Close Black Steel Pipe Nipple",
             "CHARLOTTE PIPE® 1/2 in x Close Schedule 40 Black Steel Pipe Nipple MNPT",
             [("Pipe Size", "1/2 in", "in"), ("Length", "Close", ""), ("Connection Type", "MNPT x MNPT", ""), ("Fitting Material", "Black Steel", "")])
        ]

        # Populate Input and Delivery Format sheets
        for item in gt_sample_items:
            mfg_part, desc, e1_b, uni_b, dib_b, part_mfr, cp, mfr_name, brand_name, mpn, inv_desc, mob_desc, short_desc, attrs = item
            ws_in.append([mfg_part, desc, e1_b, uni_b, dib_b, part_mfr])

            row_252 = [""] * 252
            row_252[UNILOG_252_HEADERS.index("Mfg_Part_Num")] = mfg_part
            row_252[UNILOG_252_HEADERS.index("Part_Desc")] = desc
            row_252[UNILOG_252_HEADERS.index("Part_Manuf")] = part_mfr
            row_252[UNILOG_252_HEADERS.index("MANUFACTURER_NAME")] = mfr_name
            row_252[UNILOG_252_HEADERS.index("BRAND_NAME")] = brand_name
            row_252[UNILOG_252_HEADERS.index("MANUFACTURER_PART_NUMBER")] = mpn
            row_252[UNILOG_252_HEADERS.index("Classpath")] = cp
            row_252[UNILOG_252_HEADERS.index("INVOICE_DESC")] = inv_desc
            row_252[UNILOG_252_HEADERS.index("MOBILE_DESC")] = mob_desc
            row_252[UNILOG_252_HEADERS.index("SHORT_DESC")] = short_desc
            row_252[UNILOG_252_HEADERS.index("LONG_DESC1")] = short_desc

            # Populate attributes
            for idx, (lbl, val, uom) in enumerate(attrs, start=1):
                lbl_idx = UNILOG_252_HEADERS.index(f"ATTRIBUTE_LABEL {idx}")
                val_idx = UNILOG_252_HEADERS.index(f"ATTRIBUTE_VALUE {idx}")
                uom_idx = UNILOG_252_HEADERS.index(f"ATTRIBUTE_UOM {idx}")
                row_252[lbl_idx] = lbl
                row_252[val_idx] = val
                row_252[uom_idx] = uom

            ws_out.append(row_252)

        wb_gt.save(gt_path)

        # Create Sample-1000_Items.xlsx
        wb_1000 = openpyxl.Workbook()
        ws_1000 = wb_1000.active
        ws_1000.title = "Sheet1"
        ws_1000.append(["Mfg_Part_Num", "Part_Desc", "E1_Brand", "Unilog_Brand", "DIB_Brand", "Part_Manuf"])

        # Add 1000 raw input rows combining faucets, fittings, abrasives, appliances, electrical, etc.
        bulk_raw_samples = [
            ("DCB518ASTS06G", "DCB518ASTS06G Diablo 1/2\"x18\" - Sanding Belt 6pc", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Freud Inc (2435)"),
            ("6702-000", "Moen 6702-000 Genta Single Handle Bathroom Sink Faucet Chrome 1.2 gpm 4in", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Moen Incorporated (MOEN)"),
            ("559HA-DST", "Delta 559HA-DST Trinsic Single Handle Bathroom Sink Faucet Matte Black 1.2 gpm", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Delta Faucet Company (DELTA)"),
            ("K-596-VS", "Kohler K-596-VS Simplice Pull-Down Kitchen Faucet Vibrant Stainless 1.5 gpm", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Kohler Co. (KOHLER)"),
            ("607-1/2", "Nibco 607 1/2 1/2 in Copper Wrot Pressure 90 Deg Elbow", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Nibco Inc (NIBCO)"),
            ("U008LF", "SharkBite U008LF 1/2 in Push-to-Connect Brass Coupling", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "SharkBite Plumbing Solutions (SHARK)"),
            ("NIP-050-CLOSE", "Charlotte Pipe 1/2 in x Close Black Steel Nipple MNPT", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Charlotte Pipe and Foundry (CHARL)"),
            ("PDSH4816AF", "PDSH4816AF Dishwasher SS - Display Only", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Appliance Dealers Cooperative (APPDE)"),
            ("WDTS7024RZ", "WDTS7024RZ Dishwasher SS - Display Only", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Appliance Dealers Cooperative (APPDE)"),
            ("DCB205C", "Dewalt 20V Max 5.0Ah Starter Kit", "-- Unbranded --", "-- No Unilog Brand --", "-- No DIB Brand --", "Black & Decker (2585)")
        ]

        for i in range(1000):
            sample = bulk_raw_samples[i % len(bulk_raw_samples)]
            mpn = f"{sample[0]}-{i+1}"
            desc = f"{sample[1]} Item #{i+1}"
            ws_1000.append([mpn, desc, sample[2], sample[3], sample[4], sample[5]])

        wb_1000.save(sample_1000_path)


if __name__ == '__main__':
    setup_data_files()
    print(f"Data files setup successfully in {DATA_DIR}")
